import requests
from bs4 import BeautifulSoup
from docx import Document

def extract_text_by_tags(url, selector):
    nameFromUrl = url.replace("https://", "").replace("http://", "").replace("www.", "").replace("/", "").replace(".", "_")
    response = requests.get(url)
    response.raise_for_status()  # Sprawdza, czy strona została załadowana poprawnie
 

    soup = BeautifulSoup(response.text, 'html.parser')
    elements = soup.select(selector)
    document = Document()
    document.add_heading(f'Tekst ze strony {url}', 0)
    for element in elements:
        
        if element.name in ['h1', 'h2', 'p', 'span']:  # Dodaj więcej tagów według potrzeb
            document.add_paragraph(f"Tag {element.name.upper()}:\n{element.get_text(strip=True)}\n")
            document.add_paragraph("Tłumaczenie…\n")  # Opcjonalnie, jeśli potrzebujesz miejsca na tłumaczenie

    document.save(f'{nameFromUrl}_extracted.docx')
    print('Tekst został zapisany w pliku "extracted_text.docx".')

# Przykład użycia:
url = ''  # Podmień na odpowiedni adres URL
selector = 'h1, h2, p, span'  # Zmień selektor według potrzeb
extract_text_by_tags(url, selector)
